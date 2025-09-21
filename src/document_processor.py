"""
Document processing module for extracting text from various file formats
"""
import io
import fitz  # PyMuPDF
from pypdf import PdfReader
from docx import Document
import docx2python
from pathlib import Path
import logging
from typing import Optional, Union, Dict, Any
import os

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Handles text extraction from PDF, DOCX, and TXT files
    """
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.txt'}
    
    def extract_text(self, file_path_or_bytes: Union[str, Path, bytes, io.BytesIO], 
                    file_extension: str) -> Dict[str, Any]:
        """
        Extract text from document based on file extension
        
        Args:
            file_path_or_bytes: File path, bytes, or BytesIO object
            file_extension: File extension (.pdf, .docx, .txt)
            
        Returns:
            Dict containing extracted text and metadata
        """
        try:
            file_extension = file_extension.lower()
            
            if file_extension == '.pdf':
                return self._extract_pdf_text(file_path_or_bytes)
            elif file_extension == '.docx':
                return self._extract_docx_text(file_path_or_bytes)
            elif file_extension == '.txt':
                return self._extract_txt_text(file_path_or_bytes)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            return {"text": "", "error": str(e), "pages": 0}
    
    def _extract_pdf_text(self, file_input: Union[str, Path, bytes, io.BytesIO]) -> Dict[str, Any]:
        """Extract text from PDF using PyMuPDF (primary) with pypdf fallback"""
        try:
            # Try PyMuPDF first (better for complex layouts)
            if isinstance(file_input, (str, Path)):
                doc = fitz.open(file_input)
            else:
                # Handle bytes or BytesIO
                if isinstance(file_input, io.BytesIO):
                    doc = fitz.open(stream=file_input.getvalue(), filetype="pdf")
                else:
                    doc = fitz.open(stream=file_input, filetype="pdf")
            
            full_text = ""
            pages = len(doc)
            
            for page_num in range(pages):
                page = doc[page_num]
                full_text += page.get_text() + "\n\n"
            
            doc.close()
            
            return {
                "text": full_text.strip(),
                "pages": pages,
                "method": "PyMuPDF",
                "error": None
            }
            
        except Exception as e:
            logger.warning(f"PyMuPDF failed, trying pypdf: {str(e)}")
            # Fallback to pypdf
            try:
                if isinstance(file_input, (str, Path)):
                    with open(file_input, 'rb') as file:
                        reader = PdfReader(file)
                        full_text = ""
                        for page in reader.pages:
                            full_text += page.extract_text() + "\n\n"
                else:
                    # Reset BytesIO position if needed
                    if hasattr(file_input, 'seek'):
                        file_input.seek(0)
                    reader = PdfReader(file_input)
                    full_text = ""
                    for page in reader.pages:
                        full_text += page.extract_text() + "\n\n"
                
                return {
                    "text": full_text.strip(),
                    "pages": len(reader.pages),
                    "method": "pypdf",
                    "error": None
                }
                
            except Exception as e2:
                logger.error(f"Both PDF extraction methods failed: {str(e2)}")
                return {
                    "text": "",
                    "pages": 0,
                    "method": "none",
                    "error": f"PDF extraction failed: {str(e2)}"
                }
    
    def _extract_docx_text(self, file_input: Union[str, Path, bytes, io.BytesIO]) -> Dict[str, Any]:
        """Extract text from DOCX using python-docx (primary) with docx2python fallback"""
        try:
            # Try python-docx first
            if isinstance(file_input, (str, Path)):
                doc = Document(file_input)
            else:
                # Handle bytes or BytesIO
                if isinstance(file_input, io.BytesIO):
                    doc = Document(file_input)
                else:
                    doc = Document(io.BytesIO(file_input))
            
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += cell.text + "\t"
                    full_text += "\n"
            
            return {
                "text": full_text.strip(),
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "method": "python-docx",
                "error": None
            }
            
        except Exception as e:
            logger.warning(f"python-docx failed, trying docx2python: {str(e)}")
            # Fallback to docx2python
            try:
                if isinstance(file_input, (str, Path)):
                    result = docx2python.docx2python(file_input)
                else:
                    # docx2python expects file path, so we'll save temporarily
                    import tempfile
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                        if isinstance(file_input, io.BytesIO):
                            tmp_file.write(file_input.getvalue())
                        else:
                            tmp_file.write(file_input)
                        tmp_file.flush()
                        result = docx2python.docx2python(tmp_file.name)
                    os.unlink(tmp_file.name)
                
                return {
                    "text": result.text,
                    "method": "docx2python",
                    "error": None
                }
                
            except Exception as e2:
                logger.error(f"Both DOCX extraction methods failed: {str(e2)}")
                return {
                    "text": "",
                    "method": "none", 
                    "error": f"DOCX extraction failed: {str(e2)}"
                }
    
    def _extract_txt_text(self, file_input: Union[str, Path, bytes, io.BytesIO]) -> Dict[str, Any]:
        """Extract text from TXT file"""
        try:
            if isinstance(file_input, (str, Path)):
                with open(file_input, 'r', encoding='utf-8') as file:
                    text = file.read()
            else:
                if isinstance(file_input, io.BytesIO):
                    text = file_input.getvalue().decode('utf-8')
                else:
                    text = file_input.decode('utf-8')
            
            return {
                "text": text,
                "lines": len(text.splitlines()),
                "method": "direct",
                "error": None
            }
            
        except UnicodeDecodeError:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    if isinstance(file_input, (str, Path)):
                        with open(file_input, 'r', encoding=encoding) as file:
                            text = file.read()
                    else:
                        if isinstance(file_input, io.BytesIO):
                            text = file_input.getvalue().decode(encoding)
                        else:
                            text = file_input.decode(encoding)
                    
                    return {
                        "text": text,
                        "lines": len(text.splitlines()),
                        "method": f"direct-{encoding}",
                        "error": None
                    }
                except:
                    continue
            
            return {
                "text": "",
                "method": "none",
                "error": "Could not decode text file with any encoding"
            }
        
        except Exception as e:
            logger.error(f"TXT extraction failed: {str(e)}")
            return {
                "text": "",
                "method": "none",
                "error": f"TXT extraction failed: {str(e)}"
            }
    
    def get_file_info(self, file_input: Union[str, Path, bytes, io.BytesIO], 
                     filename: str = None) -> Dict[str, Any]:
        """Get file information and validate"""
        try:
            if isinstance(file_input, (str, Path)):
                file_path = Path(file_input)
                size = file_path.stat().st_size
                extension = file_path.suffix.lower()
                name = file_path.name
            else:
                if hasattr(file_input, 'name') and file_input.name:
                    name = file_input.name
                    extension = Path(name).suffix.lower()
                elif filename:
                    name = filename
                    extension = Path(filename).suffix.lower()
                else:
                    name = "unknown"
                    extension = ""
                
                if isinstance(file_input, io.BytesIO):
                    size = len(file_input.getvalue())
                else:
                    size = len(file_input) if isinstance(file_input, bytes) else 0
            
            is_supported = extension in self.supported_formats
            
            return {
                "name": name,
                "extension": extension,
                "size": size,
                "size_mb": round(size / (1024 * 1024), 2),
                "is_supported": is_supported
            }
            
        except Exception as e:
            return {
                "name": filename or "unknown",
                "extension": "",
                "size": 0,
                "size_mb": 0,
                "is_supported": False,
                "error": str(e)
            }